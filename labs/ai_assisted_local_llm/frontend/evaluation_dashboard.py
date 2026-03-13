"""Streamlit Evaluation Dashboard — multidimensional transcript evaluation."""

from __future__ import annotations

import io
import json
import os
import re
from pathlib import Path

import plotly.graph_objects as go
import requests
import streamlit as st
from docx import Document

st.set_page_config(page_title="Evaluation Dashboard", layout="wide")

API_BASE = os.getenv("API_BASE_URL", "http://localhost:8000")
SAMPLE_DIR = Path(__file__).resolve().parent.parent / "sample_data"
PREFILL_DIR = Path(__file__).resolve().parent / ".eval_prefill"
SHARED_PREFILL_FILE = PREFILL_DIR / "latest_eval_payload.json"


# ── File parsing helpers ────────────────────────────────────────────────────


def _parse_case_from_docx(file_bytes: bytes) -> dict:
    """Extract a case description dict from a Word document.

    Expects the document to contain structured sections with headings like
    "Chief Complaint", "HPI", "Past Medical History", etc. Falls back to
    treating the entire document body as the HPI if no headings are found.
    """
    doc = Document(io.BytesIO(file_bytes))

    sections: dict[str, list[str]] = {}
    current_heading = "_preamble"
    sections[current_heading] = []

    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            current_heading = para.text.strip().lower()
            sections[current_heading] = []
        elif para.text.strip():
            sections[current_heading].append(para.text.strip())

    def _get(keys: list[str]) -> str:
        for k in keys:
            if k in sections and sections[k]:
                return "\n".join(sections[k])
        return ""

    def _get_list(keys: list[str]) -> list[str]:
        text = _get(keys)
        if not text:
            return []
        return [line.lstrip("•-– ").strip() for line in text.split("\n") if line.strip()]

    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if len(sections) <= 1:
        return {
            "chief_complaint": "",
            "hpi": full_text,
            "emotional_presentation": "",
        }

    demo: dict = {}
    demo_text = _get(["demographics", "patient demographics", "_preamble"])
    age_match = re.search(r"(\d{1,3})\s*(?:year|yo|y/?o)", demo_text, re.IGNORECASE)
    sex_match = re.search(r"\b(male|female|man|woman)\b", demo_text, re.IGNORECASE)
    if age_match:
        demo["age"] = int(age_match.group(1))
    if sex_match:
        demo["sex"] = sex_match.group(1).lower()

    return {
        "demographics": demo,
        "chief_complaint": _get(["chief complaint", "cc"]),
        "hpi": _get(["hpi", "history of present illness", "history"]),
        "pmh": _get_list(["past medical history", "pmh", "medical history"]),
        "medications": _get_list(["medications", "meds", "current medications"]),
        "allergies": _get_list(["allergies"]),
        "social_history": {"narrative": _get(["social history", "social hx"])},
        "family_history": _get_list(["family history", "family hx", "fhx"]),
        "ros": {"narrative": _get(["review of systems", "ros"])},
        "physical_exam_findings": {"narrative": _get(["physical exam", "physical examination", "pe"])},
        "labs": {"narrative": _get(["labs", "laboratory", "lab results"])},
        "imaging": _get_list(["imaging", "radiology"]),
        "differential_diagnosis": _get_list(["differential", "differential diagnosis", "ddx"]),
        "final_diagnosis": _get(["final diagnosis", "diagnosis", "assessment"]),
        "emotional_presentation": _get(["emotional presentation", "affect", "patient demeanor"]),
    }


def _parse_transcript_from_docx(file_bytes: bytes) -> dict:
    """Extract a transcript dict from a Word document.

    Expects lines formatted as "Student: ..." or "Patient: ...".
    """
    doc = Document(io.BytesIO(file_bytes))
    turns: list[dict] = []
    turn_num = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        match = re.match(r"^(Student|Patient)\s*:\s*(.+)", text, re.IGNORECASE)
        if match:
            turn_num += 1
            turns.append({
                "turn_number": turn_num,
                "speaker": match.group(1).capitalize(),
                "content": match.group(2).strip(),
            })

    return {"turns": turns}


def _parse_transcript_from_text(text: str) -> dict:
    """Parse a plain-text transcript with 'Student:' / 'Patient:' prefixes."""
    turns: list[dict] = []
    turn_num = 0

    for line in text.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        match = re.match(r"^(Student|Patient)\s*:\s*(.+)", line, re.IGNORECASE)
        if match:
            turn_num += 1
            turns.append({
                "turn_number": turn_num,
                "speaker": match.group(1).capitalize(),
                "content": match.group(2).strip(),
            })

    return {"turns": turns}


def _read_uploaded_file(uploaded_file, file_type: str) -> str | None:
    """Read an uploaded file (JSON or DOCX) and return JSON string, or None."""
    if uploaded_file is None:
        return None

    raw = uploaded_file.read()
    name = uploaded_file.name.lower()

    if name.endswith(".json"):
        return raw.decode("utf-8")

    if name.endswith(".docx"):
        if file_type == "case":
            parsed = _parse_case_from_docx(raw)
        else:
            parsed = _parse_transcript_from_docx(raw)
        return json.dumps(parsed, indent=2)

    return None


def _load_prefill_from_shared_file() -> None:
    """
    Auto-load prefill data from a shared local JSON file written by interview.py.
    No query param is needed.
    """
    if not SHARED_PREFILL_FILE.exists():
        return

    try:
        payload = json.loads(SHARED_PREFILL_FILE.read_text(encoding="utf-8"))
    except Exception as e:
        st.error(f"Failed to read shared prefill file: {e}")
        return

    meta = payload.get("meta", {})
    signature = (
        str(meta.get("conversation_id", "")),
        str(meta.get("created_at", "")),
        str(meta.get("case_number", "")),
    )

    if st.session_state.get("_loaded_prefill_signature") == signature:
        return

    try:
        st.session_state["case_input"] = json.dumps(payload.get("case_input", {}), indent=2)
        st.session_state["transcript_input"] = json.dumps(payload.get("transcript_input", {}), indent=2)
        st.session_state["_loaded_prefill_signature"] = signature
        st.session_state["eval_prefilled_from_interview"] = True
    except Exception as e:
        st.error(f"Failed to load prefill data into session: {e}")


# ── API helpers ──────────────────────────────────────────────────────────────

def _api_evaluate(payload: dict) -> dict:
    connect_timeout = 10
    read_timeout = 600  # allow long evaluation runs
    resp = requests.post(
        f"{API_BASE}/api/v1/evaluate/",
        json=payload,
        timeout=(connect_timeout, read_timeout),
    )
    print(f"API response status: {resp.status_code}")
    resp.raise_for_status()
    return resp.json()


# ── Chart helpers ────────────────────────────────────────────────────────────


def _radar_chart(dimensions: list[dict], title: str) -> go.Figure:
    """Create a radar/spider chart from dimension scores."""
    names = [d["dimension"] for d in dimensions]
    scores = [d["score"] for d in dimensions]
    names_closed = names + [names[0]]
    scores_closed = scores + [scores[0]]

    fig = go.Figure()
    fig.add_trace(
        go.Scatterpolar(
            r=scores_closed,
            theta=names_closed,
            fill="toself",
            name=title,
            line=dict(color="#636EFA"),
        )
    )
    fig.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 5])),
        title=title,
        showlegend=False,
        height=400,
    )
    return fig


def _render_score_badge(score: int) -> str:
    if score >= 4:
        color = "green"
    elif score >= 3:
        color = "orange"
    else:
        color = "red"
    return f":{color}[**{score}/5**]"


# ── Load sample data ────────────────────────────────────────────────────────


def _load_sample_case() -> dict | None:
    path = SAMPLE_DIR / "sample_case_chest_pain.json"
    if path.exists():
        return json.loads(path.read_text())
    return None


def _load_sample_transcript(quality: str) -> dict | None:
    path = SAMPLE_DIR / f"sample_transcript_{quality}.json"
    if path.exists():
        return json.loads(path.read_text())
    return None


# ── Sidebar ──────────────────────────────────────────────────────────────────

st.sidebar.title("Evaluation Settings")

model_choice = st.sidebar.selectbox(
    "LLM Model",
    ["gpt-4o","llama3.2","gemma3"],
    index=0,
)

eval_layers = st.sidebar.multiselect(
    "Evaluation Layers",
    ["Case Fidelity", "Student Performance"],
    default=["Case Fidelity", "Student Performance"],
)

layer_map = {
    frozenset(["Case Fidelity", "Student Performance"]): "both",
    frozenset(["Case Fidelity"]): "case_fidelity",
    frozenset(["Student Performance"]): "student_performance",
}
selected_layer = layer_map.get(frozenset(eval_layers), "both")

st.sidebar.divider()
st.sidebar.caption("Sample data available for demo")
if st.sidebar.button("Load Sample Case"):
    sample = _load_sample_case()
    if sample:
        st.session_state["case_input"] = json.dumps(sample, indent=2)
        st.rerun()

sample_quality = st.sidebar.radio("Sample transcript quality", ["good", "poor"], horizontal=True)
if st.sidebar.button("Load Sample Transcript"):
    sample = _load_sample_transcript(sample_quality)
    if sample:
        st.session_state["transcript_input"] = json.dumps(sample, indent=2)
        st.rerun()

# Auto-load from shared file before rendering widgets
_load_prefill_from_shared_file()

# ── Main Area ────────────────────────────────────────────────────────────────

st.title("Multidimensional Evaluation Dashboard")

if st.session_state.get("eval_prefilled_from_interview"):
    st.success("Case description and transcript were preloaded from the interview app.")

if "eval_running" not in st.session_state:
    st.session_state["eval_running"] = False

tab_run, tab_batch, tab_analytics = st.tabs(["Run Evaluation", "Batch Evaluation", "Analytics"])

# ── Tab 1: Run Evaluation ────────────────────────────────────────────────────

with tab_run:
    st.header("Evaluate a Transcript")

    col_case, col_transcript = st.columns(2)

    with col_case:
        st.subheader("Case Description")
        case_file = st.file_uploader(
            "Upload case file", type=["json", "docx"], key="case_file",
        )
        if case_file:
            file_id = f"{case_file.name}_{case_file.size}"
            if st.session_state.get("_last_case_file_id") != file_id:
                parsed = _read_uploaded_file(case_file, "case")
                if parsed:
                    st.session_state["case_input"] = parsed
                    st.session_state["_last_case_file_id"] = file_id
                    st.rerun()

        case_input = st.text_area(
            "Or paste case JSON",
            height=300,
            key="case_input",
            help="JSON object with demographics, chief_complaint, hpi, pmh, etc.",
        )

    with col_transcript:
        st.subheader("Transcript")
        transcript_file = st.file_uploader(
            "Upload transcript file", type=["json", "docx"], key="transcript_file",
        )
        if transcript_file:
            file_id = f"{transcript_file.name}_{transcript_file.size}"
            if st.session_state.get("_last_transcript_file_id") != file_id:
                parsed = _read_uploaded_file(transcript_file, "transcript")
                if parsed:
                    st.session_state["transcript_input"] = parsed
                    st.session_state["_last_transcript_file_id"] = file_id
                    st.rerun()

        transcript_input = st.text_area(
            "Or paste transcript JSON / plain text",
            height=300,
            key="transcript_input",
            help='JSON with "turns" array, or plain text with "Student:" / "Patient:" prefixes',
        )

    with st.expander("Preview inputs", expanded=bool(case_input or transcript_input)):
        prev_col1, prev_col2 = st.columns(2)
        with prev_col1:
            if case_input:
                try:
                    parsed_case = json.loads(case_input)
                    st.markdown(f"**Chief Complaint:** {parsed_case.get('chief_complaint', 'N/A')}")
                    st.markdown(f"**Final Diagnosis:** {parsed_case.get('final_diagnosis', 'N/A')}")
                    demo = parsed_case.get("demographics", {})
                    if demo:
                        st.markdown(
                            f"**Patient:** {demo.get('age', '?')}yo {demo.get('sex', '?')}"
                        )
                except json.JSONDecodeError:
                    st.warning("Invalid JSON in case input")
        with prev_col2:
            if transcript_input:
                try:
                    parsed_transcript = json.loads(transcript_input)
                    turns = parsed_transcript.get("turns", [])
                    st.markdown(f"**Turns:** {len(turns)}")
                    if turns:
                        first = turns[0]
                        st.markdown(
                            f"**First turn:** {first.get('speaker', '?')}: "
                            f"{first.get('content', '')[:100]}..."
                        )
                except json.JSONDecodeError:
                    parsed_transcript = _parse_transcript_from_text(transcript_input)
                    if parsed_transcript["turns"]:
                        st.markdown(f"**Turns (parsed from text):** {len(parsed_transcript['turns'])}")
                    else:
                        st.warning("Could not parse transcript. Use JSON or 'Student:'/'Patient:' format.")

    run_eval_clicked = st.button(
        "Run Evaluation",
        type="primary",
        use_container_width=True,
        disabled=st.session_state["eval_running"],
    )

    if run_eval_clicked:
        st.session_state["eval_running"] = True
        st.rerun()

    if st.session_state["eval_running"]:
        if not case_input or not transcript_input:
            st.error("Please provide both a case description and a transcript.")
            st.session_state["eval_running"] = False
        else:
            try:
                case_data = json.loads(case_input)
            except json.JSONDecodeError:
                st.error("Invalid JSON in case input. Upload a JSON or DOCX file.")
                st.session_state["eval_running"] = False
                st.stop()

            try:
                transcript_data = json.loads(transcript_input)
            except json.JSONDecodeError:
                transcript_data = _parse_transcript_from_text(transcript_input)
                if not transcript_data["turns"]:
                    st.error(
                        "Could not parse transcript. Provide JSON or plain text "
                        "with 'Student:' / 'Patient:' prefixes."
                    )
                    st.session_state["eval_running"] = False
                    st.stop()

            payload = {
                "case_description": case_data,
                "transcript": transcript_data,
                "layer": selected_layer,
                "model": model_choice,
            }

            print(f"Sending evaluation request with payload: {json.dumps(payload, indent=2)}")

            try:
                with st.spinner("Evaluating transcript... this may take few minutes per layer."):
                    result = _api_evaluate(payload)

                st.session_state["eval_result"] = result
                st.success("Evaluation complete!")

            except requests.HTTPError as e:
                st.error(f"API error: {e.response.status_code} — {e.response.text}")
            except requests.ConnectionError:
                st.error("Could not connect to the API. Is the FastAPI server running?")
            finally:
                st.session_state["eval_running"] = False

    if "eval_result" in st.session_state:
        result = st.session_state["eval_result"]

        st.divider()
        st.subheader("Evaluation Results")

        meta_cols = st.columns(3)
        meta_cols[0].metric("Model Used", result.get("model_used", "N/A"))
        tokens = result.get("token_usage", {})
        meta_cols[1].metric("Input Tokens", f"{tokens.get('input_tokens', 0):,}")
        meta_cols[2].metric("Output Tokens", f"{tokens.get('output_tokens', 0):,}")

        for eval_result in result.get("results", []):
            layer_name = eval_result["layer"].replace("_", " ").title()
            st.markdown(f"### {layer_name}")

            st.metric(
                f"{layer_name} — Weighted Score",
                f"{eval_result['weighted_total']:.2f} / 5.00",
            )

            st.plotly_chart(
                _radar_chart(eval_result["dimensions"], layer_name),
                use_container_width=True,
            )

            for dim in eval_result["dimensions"]:
                with st.expander(
                    f"{dim['dimension']} — {_render_score_badge(dim['score'])} (weight: {dim['weight']})",
                    expanded=False,
                ):
                    st.markdown(f"**Rationale:** {dim['rationale']}")

                    if dim.get("strengths"):
                        st.markdown("**Strengths:**")
                        for s in dim["strengths"]:
                            st.markdown(f"- {s}")

                    if dim.get("growth_areas"):
                        st.markdown("**Growth Areas:**")
                        for g in dim["growth_areas"]:
                            st.markdown(f"- {g}")

                    if dim.get("evidence"):
                        st.markdown("**Evidence Citations:**")
                        for ev in dim["evidence"]:
                            st.info(
                                f"**Turn {ev['turn_number']}** ({ev['speaker']}): "
                                f'"{ev["quote"]}"\n\n*{ev["relevance"]}*'
                            )

            st.markdown("---")
            st.markdown(f"**Overall Summary:** {eval_result['overall_summary']}")
            st.success(f"**Top Recommendation:** {eval_result['top_recommendation']}")

        st.divider()
        st.download_button(
            "Download Evaluation JSON",
            data=json.dumps(result, indent=2, default=str),
            file_name="evaluation_result.json",
            mime="application/json",
        )

# ── Tab 2: Batch Evaluation ─────────────────────────────────────────────────

with tab_batch:
    st.header("Batch Evaluation")
    st.info("Upload a case and multiple transcripts to evaluate them all at once.")

    batch_case_file = st.file_uploader(
        "Upload case file (shared across all transcripts)",
        type=["json", "docx"],
        key="batch_case_file",
    )
    batch_case_input = st.text_area(
        "Or paste case JSON",
        height=200,
        key="batch_case",
    )

    if batch_case_file:
        parsed = _read_uploaded_file(batch_case_file, "case")
        if parsed:
            batch_case_input = parsed

    batch_files = st.file_uploader(
        "Upload transcript files",
        type=["json", "docx"],
        accept_multiple_files=True,
        key="batch_transcripts",
    )

    if st.button("Run Batch Evaluation", type="primary"):
        if not batch_case_input or not batch_files:
            st.error("Please provide a case and at least one transcript file.")
        else:
            try:
                case_data = json.loads(batch_case_input)
            except json.JSONDecodeError:
                st.error("Invalid case JSON.")
                st.stop()

            progress = st.progress(0, text="Evaluating transcripts...")
            results = []
            for i, f in enumerate(batch_files):
                parsed_t = _read_uploaded_file(f, "transcript")
                if not parsed_t:
                    st.warning(f"Could not parse {f.name}, skipping.")
                    continue
                transcript_data = json.loads(parsed_t)

                payload = {
                    "case_description": case_data,
                    "transcript": transcript_data,
                    "layer": selected_layer,
                    "model": model_choice,
                }
                try:
                    resp = _api_evaluate(payload)
                    results.append(resp)
                except Exception as e:
                    st.warning(f"Evaluation failed for {f.name}: {e}")

                progress.progress(
                    (i + 1) / len(batch_files),
                    text=f"Evaluated {i + 1}/{len(batch_files)} transcripts",
                )

            st.session_state["batch_results"] = results
            st.success(f"Batch evaluation complete — {len(results)} transcripts evaluated.")

    if "batch_results" in st.session_state:
        batch = st.session_state["batch_results"]

        rows = []
        for i, resp in enumerate(batch):
            row = {"Transcript": i + 1}
            for er in resp.get("results", []):
                label = er["layer"].replace("_", " ").title()
                row[label] = er["weighted_total"]
            rows.append(row)

        st.dataframe(rows, use_container_width=True)

        import plotly.express as px

        all_scores = []
        for resp in batch:
            for er in resp.get("results", []):
                for dim in er["dimensions"]:
                    all_scores.append({
                        "Layer": er["layer"].replace("_", " ").title(),
                        "Dimension": dim["dimension"],
                        "Score": dim["score"],
                    })

        if all_scores:
            fig = px.box(
                all_scores,
                x="Dimension",
                y="Score",
                color="Layer",
                title="Score Distribution Across Batch",
            )
            fig.update_layout(yaxis=dict(range=[0, 5.5]))
            st.plotly_chart(fig, use_container_width=True)

# ── Tab 3: Analytics ─────────────────────────────────────────────────────────

with tab_analytics:
    st.header("Evaluation Analytics")
    st.info(
        "Analytics will show trends from stored evaluations. "
        "Run evaluations first to populate this view."
    )

    if "batch_results" in st.session_state and st.session_state["batch_results"]:
        batch = st.session_state["batch_results"]
        import pandas as pd
        import plotly.express as px

        heatmap_data = []
        for i, resp in enumerate(batch):
            for er in resp.get("results", []):
                for dim in er["dimensions"]:
                    heatmap_data.append({
                        "Transcript": f"T{i + 1}",
                        "Dimension": dim["dimension"],
                        "Score": dim["score"],
                    })

        if heatmap_data:
            df = pd.DataFrame(heatmap_data)
            pivot = df.pivot_table(
                index="Dimension", columns="Transcript", values="Score", aggfunc="mean"
            )

            fig = px.imshow(
                pivot,
                text_auto=True,
                color_continuous_scale="RdYlGn",
                zmin=1,
                zmax=5,
                title="Dimension x Transcript Heatmap",
            )
            st.plotly_chart(fig, use_container_width=True)

            fig2 = px.histogram(
                df,
                x="Score",
                color="Dimension",
                barmode="overlay",
                title="Score Distribution by Dimension",
                nbins=5,
            )
            st.plotly_chart(fig2, use_container_width=True)

            csv = df.to_csv(index=False)
            st.download_button(
                "Export Analytics CSV",
                data=csv,
                file_name="evaluation_analytics.csv",
                mime="text/csv",
            )
    else:
        st.caption("No evaluation data available yet. Run evaluations in the other tabs first.")