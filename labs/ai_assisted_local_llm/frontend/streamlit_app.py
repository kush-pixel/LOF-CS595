"""Streamlit frontend for the Medical Case Generator."""

from __future__ import annotations

import io
import json
import os
from datetime import datetime, timezone

from docx import Document
import streamlit as st

from api_client import (
    delete_case,
    generate_case,
    get_case,
    get_case_by_number,
    list_cases,
    patch_case,
)

st.set_page_config(page_title="Medical Case Generator", layout="wide")

SPECIALTIES = [
    "",
    "cardiology",
    "pulmonology",
    "neurology",
    "gastroenterology",
    "endocrinology",
    "nephrology",
    "infectious disease",
    "hematology/oncology",
    "rheumatology",
    "psychiatry",
    "pediatrics",
    "surgery",
    "emergency medicine",
    "internal medicine",
]

TAB_SEARCH = "Search & Browse"
TAB_GENERATE = "Generate"
TAB_DETAIL = "Case Detail"


# ── Helper: build case Word doc ──────────────────────────────────────────────
def _build_case_docx(case: dict) -> bytes:
    """Build a Word document from a case dict and return raw bytes."""
    doc = Document()
    doc.add_heading(case.get("case_title", "Untitled Case"), level=1)
    doc.add_paragraph(
        f"Case #{case.get('case_number', '—')} | "
        f"Specialty: {case.get('specialty', '')} | "
        f"Difficulty: {case.get('difficulty', '')} | "
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )

    # Demographics
    demo = case.get("demographics") or {}
    if demo:
        doc.add_heading("Demographics", level=2)
        for k, v in demo.items():
            if v is not None:
                doc.add_paragraph(f"{k.replace('_', ' ').title()}: {v}")

    # Chief Complaint & HPI
    hpi = case.get("chief_complaint_hpi") or {}
    if hpi:
        doc.add_heading("Chief Complaint", level=2)
        doc.add_paragraph(hpi.get("chief_complaint", ""))
        doc.add_heading("HPI", level=2)
        doc.add_paragraph(hpi.get("hpi_narrative", ""))

    # Past Medical History
    pmh = case.get("past_medical_history") or {}
    if pmh:
        doc.add_heading("Past Medical History", level=2)
        for cond in pmh.get("conditions", []):
            doc.add_paragraph(cond, style="List Bullet")

    # Medications
    meds = case.get("medications") or []
    if meds:
        doc.add_heading("Medications", level=2)
        for m in meds:
            doc.add_paragraph(
                f"{m.get('name', '?')} {m.get('dose', '')} {m.get('route', '')} {m.get('frequency', '')}".strip(),
                style="List Bullet",
            )

    # Allergies
    allergies = case.get("allergies") or []
    if allergies:
        doc.add_heading("Allergies", level=2)
        for a in allergies:
            doc.add_paragraph(
                f"{a.get('substance', '?')} — {a.get('reaction', '')} ({a.get('severity', '')})",
                style="List Bullet",
            )

    # Social History
    social = case.get("social_history") or {}
    if social:
        doc.add_heading("Social History", level=2)
        for k, v in social.items():
            if v is not None:
                doc.add_paragraph(f"{k.replace('_', ' ').title()}: {v}")

    # Family History
    fhx = case.get("family_history") or []
    if fhx:
        doc.add_heading("Family History", level=2)
        for member in fhx:
            if isinstance(member, dict):
                doc.add_paragraph(
                    f"{member.get('relation', '?')}: {', '.join(member.get('conditions', []))}",
                    style="List Bullet",
                )
            else:
                doc.add_paragraph(str(member), style="List Bullet")

    # Physical Exam
    pe = case.get("physical_exam") or {}
    if pe:
        doc.add_heading("Physical Exam", level=2)
        _write_nested_dict(doc, pe)

    # Diagnostics
    diag = case.get("diagnostics") or {}
    if diag:
        doc.add_heading("Diagnostics", level=2)
        _write_nested_dict(doc, diag)

    # Assessment
    assess = case.get("assessment") or {}
    if assess:
        doc.add_heading("Assessment", level=2)
        if assess.get("working_diagnosis"):
            doc.add_paragraph(f"Working Diagnosis: {assess['working_diagnosis']}")
        if assess.get("final_diagnosis"):
            doc.add_paragraph(f"Final Diagnosis: {assess['final_diagnosis']}")
        for ddx in assess.get("differential_diagnoses", []):
            if isinstance(ddx, dict):
                doc.add_paragraph(
                    f"{ddx.get('rank', '?')}. {ddx.get('diagnosis', '?')} — {ddx.get('reasoning', '')}",
                    style="List Bullet",
                )

    # Plan
    plan = case.get("plan") or {}
    if plan:
        doc.add_heading("Plan", level=2)
        for step in plan.get("steps", []):
            if isinstance(step, dict):
                doc.add_paragraph(
                    f"[{step.get('category', '')}] {step.get('description', '')} (Priority: {step.get('priority', '')})",
                    style="List Bullet",
                )
        if plan.get("disposition"):
            doc.add_paragraph(f"Disposition: {plan['disposition']}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _write_nested_dict(doc: Document, data: dict, level: int = 0) -> None:
    """Recursively write a nested dict into a Word document."""
    for k, v in data.items():
        if v is None:
            continue
        label = k.replace("_", " ").title()
        if isinstance(v, dict):
            doc.add_paragraph(f"{label}:")
            _write_nested_dict(doc, v, level + 1)
        elif isinstance(v, list):
            doc.add_paragraph(f"{label}:")
            for item in v:
                if isinstance(item, dict):
                    doc.add_paragraph(json.dumps(item), style="List Bullet")
                else:
                    doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph(f"{label}: {v}")


# ── Helper: load a case into session state for detail view ────────────────────
def _load_case(case: dict, switch_to_detail: bool = True) -> None:
    st.session_state["detail_case"] = case
    if switch_to_detail:
        st.session_state["active_tab"] = TAB_DETAIL


def _render_field(label: str, value, key_prefix: str, editing: bool):
    """Render a single field as read-only text or an editable input."""
    if editing:
        if isinstance(value, bool):
            return st.checkbox(label, value=value, key=f"{key_prefix}_{label}")
        if isinstance(value, int):
            return st.number_input(label, value=value, key=f"{key_prefix}_{label}")
        if isinstance(value, float):
            return st.number_input(label, value=value, format="%.2f", key=f"{key_prefix}_{label}")
        if isinstance(value, list):
            text = st.text_area(label, value=json.dumps(value, indent=2), key=f"{key_prefix}_{label}")
            try:
                return json.loads(text)
            except json.JSONDecodeError:
                return value
        if value is None:
            return st.text_input(label, value="", key=f"{key_prefix}_{label}") or None
        return st.text_area(label, value=str(value), key=f"{key_prefix}_{label}")
    else:
        if value is not None and value != "" and value != []:
            st.markdown(f"**{label}:** {value}")
        return value


def _render_dict_section(title: str, data: dict | None, key_prefix: str, editing: bool) -> dict | None:
    """Render a dict section (like demographics, vitals) with per-field editing."""
    if data is None and not editing:
        return None
    if data is None:
        data = {}
    result = {}
    for field_name, field_value in data.items():
        result[field_name] = _render_field(field_name.replace("_", " ").title(), field_value, key_prefix, editing)
    return result


# ── Session-state based tab navigation ────────────────────────────────────────
if "active_tab" not in st.session_state:
    st.session_state["active_tab"] = TAB_SEARCH

if "is_generating_case" not in st.session_state:
    st.session_state["is_generating_case"] = False

if "previous_tab" not in st.session_state:
    st.session_state["previous_tab"] = st.session_state["active_tab"]

selected_tab = st.radio(
    "Navigation",
    [TAB_SEARCH, TAB_GENERATE, TAB_DETAIL],
    index=[TAB_SEARCH, TAB_GENERATE, TAB_DETAIL].index(st.session_state["active_tab"]),
    horizontal=True,
    label_visibility="collapsed",
)

previous_tab = st.session_state.get("active_tab", TAB_SEARCH)
st.session_state["active_tab"] = selected_tab

# When user newly switches into Generate tab, reset it like a fresh page
if previous_tab != TAB_GENERATE and selected_tab == TAB_GENERATE:
    st.session_state.pop("generated_case", None)
    st.session_state["gen_spec"] = ""
    st.session_state["gen_diff"] = ""
    st.session_state["gen_prompt"] = ""
    st.session_state["ollama_model_generate"] = "openai-remote"
    st.session_state["is_generating_case"] = False

st.session_state["previous_tab"] = selected_tab


# ── Tab 1: Search & Browse ────────────────────────────────────────────────────
if selected_tab == TAB_SEARCH:
    st.header("Search & Browse Cases")

    col_search, col_number = st.columns([3, 1])
    with col_search:
        search_text = st.text_input("Search by title", placeholder="e.g. chest pain")
    with col_number:
        case_num = st.number_input("Lookup by case #", min_value=0, value=0, step=1)

    if case_num > 0:
        try:
            found = get_case_by_number(case_num)
            st.success(f"Found: {found.get('case_title', 'Untitled')}")
            if st.button("View this case", key="view_by_num"):
                _load_case(found, switch_to_detail=True)
                st.rerun()
        except Exception as e:
            st.warning(f"No case found with number {case_num}: {e}")
    else:
        col_spec, col_page = st.columns([2, 1])
        with col_spec:
            filter_specialty = st.selectbox("Filter by specialty", SPECIALTIES, key="browse_spec")
        with col_page:
            browse_page = st.number_input("Page", min_value=1, value=1, key="browse_page")

        try:
            data = list_cases(
                page=browse_page,
                specialty=filter_specialty or None,
                search=search_text or None,
            )
            st.caption(f"Page {data['page']} — {data['total']} total cases")

            for item in data.get("items", []):
                with st.container():
                    cols = st.columns([0.5, 3, 1, 1, 1])
                    cols[0].write(f"**#{item.get('case_number', '—')}**")
                    cols[1].write(f"**{item.get('case_title', 'Untitled')}**")
                    cols[2].write(item.get("specialty", ""))
                    cols[3].write(item.get("difficulty", ""))
                    if cols[4].button("View", key=f"view_{item['case_id']}"):
                        full_case = get_case(item["case_id"])
                        _load_case(full_case, switch_to_detail=True)
                        st.rerun()
        except Exception as e:
            st.error(f"Failed to load cases: {e}")


# ── Tab 2: Generate ──────────────────────────────────────────────────────────
elif selected_tab == TAB_GENERATE:
    st.header("Generate a Medical Case")

    col1, col2 = st.columns(2)
    with col1:
        gen_specialty = st.selectbox("Specialty", SPECIALTIES, key="gen_spec")
        gen_difficulty = st.selectbox("Difficulty", ["", "easy", "medium", "hard"], key="gen_diff")
    with col2:
        gen_prompt = st.text_area("Additional prompt / context", height=120, key="gen_prompt")

    DEFAULT_OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o")
    DEFAULT_OPENAI_AUDIO_MODEL = os.getenv("OPENAI_AUDIO_MODEL", "gpt-4o-audio-preview")
    OLLAMA_MODELS = ["llama3.2", "gemma3"]

    col_llm1, col_llm2 = st.columns([1, 2])

    with col_llm1:
        llm_user_provider = st.text_input(
            "LLM Provider",
            value="ollama",
            key="llm_provider_generate",
            disabled=True,
        )

    with col_llm2:
        llm_user_model = st.selectbox(
            "Select Model",
            options=["openai-remote", "llama3.2", "gemma3"],
            index=0,
            key="ollama_model_generate",
        )

    generate_clicked = st.button(
        "Generate Case",
        type="primary",
        disabled=st.session_state["is_generating_case"],
    )

    if generate_clicked:
        st.session_state["is_generating_case"] = True
        st.rerun()

    if st.session_state["is_generating_case"]:
        with st.spinner("Generating case via LLM..."):
            try:
                if llm_user_provider == "ollama" and llm_user_model == "openai-remote":
                    llm_user_provider = "openai"
                    llm_user_model = DEFAULT_OPENAI_MODEL

                case = generate_case(
                    specialty=gen_specialty or None,
                    prompt=gen_prompt or None,
                    difficulty=gen_difficulty or None,
                    llm_provider=llm_user_provider,
                    llm_model=llm_user_model,
                )
                st.session_state["generated_case"] = case
                st.success(f"Case generated: {case.get('case_title', case['case_id'])}")
            except Exception as e:
                st.error(f"Generation failed: {e}")
            finally:
                st.session_state["is_generating_case"] = False

    if "generated_case" in st.session_state:
        case = st.session_state["generated_case"]
        st.subheader(case.get("case_title", "Untitled Case"))

        with st.expander("Demographics & Vitals", expanded=True):
            st.json(case.get("demographics"))
            st.json(case.get("vitals"))
        with st.expander("Chief Complaint & HPI"):
            st.json(case.get("chief_complaint_hpi"))
        with st.expander("History"):
            st.json(case.get("past_medical_history"))
            st.json(case.get("social_history"))
            st.json(case.get("medications"))
            st.json(case.get("allergies"))
        with st.expander("Physical Exam"):
            st.json(case.get("physical_exam"))
        with st.expander("Diagnostics"):
            st.json(case.get("diagnostics"))
        with st.expander("Assessment & Plan"):
            st.json(case.get("assessment"))
            st.json(case.get("plan"))

        col_open, col_dl = st.columns(2)
        with col_open:
            if st.button("Open in Detail tab"):
                _load_case(case, switch_to_detail=True)
                st.rerun()
        with col_dl:
            docx_bytes = _build_case_docx(case)
            st.download_button(
                "Download Case (.docx)",
                data=docx_bytes,
                file_name=f"case_{case.get('case_number', case.get('case_id', 'unknown'))}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_gen_docx",
            )


# ── Tab 3: Case Detail ───────────────────────────────────────────────────────
elif selected_tab == TAB_DETAIL:
    st.header("Case Detail") 

    with st.form("load_case_form", clear_on_submit=False):
        manual_id = st.text_input("Enter Case ID to load", value="", key="manual_case_id")
        load_case_clicked = st.form_submit_button("Load case")

    if load_case_clicked:
        cleaned_case_id = manual_id.strip()

        if not cleaned_case_id:
            st.warning("Please enter a Case ID.")
        else:
            try:
                loaded = get_case(cleaned_case_id)
                st.session_state["detail_case"] = loaded
                st.session_state["active_tab"] = TAB_DETAIL
                st.success(f"Loaded case: {loaded.get('case_title', 'Untitled Case')}")
                st.rerun()
            except Exception as e:
                st.error(f"Failed to load case: {e}")

    case = st.session_state.get("detail_case")
    if not case:
        st.info("Select a case from Search & Browse or Generate tab, or enter a Case ID above.")
    else:
        case_id = case["case_id"]

        # Header
        st.subheader(case.get("case_title", "Untitled Case"))
        st.caption(
            f"Case #{case.get('case_number', '—')} | "
            f"ID: {case_id} | "
            f"Specialty: {case.get('specialty', '')} | "
            f"Difficulty: {case.get('difficulty', '')}"
        )

        # Download as Word doc
        docx_bytes = _build_case_docx(case)
        st.download_button(
            "Download Case (.docx)",
            data=docx_bytes,
            file_name=f"case_{case.get('case_number', case_id)}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_detail_docx",
        )

        # ── Section: Demographics ─────────────────────────────────────
        SECTIONS = [
            ("Demographics", "demographics"),
            ("Vitals", "vitals"),
            ("Chief Complaint & HPI", "chief_complaint_hpi"),
            ("Past Medical History", "past_medical_history"),
            ("Social History", "social_history"),
            ("Physical Exam", "physical_exam"),
            ("Assessment", "assessment"),
        ]

        for section_label, section_key in SECTIONS:
            section_data = case.get(section_key)
            with st.expander(section_label, expanded=(section_key == "demographics")):
                editing = st.toggle(f"Edit {section_label}", key=f"edit_{section_key}")

                if isinstance(section_data, dict):
                    updated = _render_dict_section(section_label, section_data, section_key, editing)
                else:
                    if editing:
                        raw = st.text_area(
                            f"{section_label} JSON",
                            value=json.dumps(section_data, indent=2, default=str) if section_data else "{}",
                            key=f"raw_{section_key}",
                            height=200,
                        )
                        try:
                            updated = json.loads(raw)
                        except json.JSONDecodeError:
                            st.warning("Invalid JSON")
                            updated = section_data
                    else:
                        st.json(section_data)
                        updated = section_data

                if editing and st.button(f"Save {section_label}", key=f"save_{section_key}"):
                    try:
                        result = patch_case(case_id, {"case_data": {section_key: updated}})
                        st.success(f"{section_label} saved!")
                        _load_case(result, switch_to_detail=True)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Save failed: {e}")

        # ── Section: Medications (list) ───────────────────────────────
        with st.expander("Medications"):
            editing_meds = st.toggle("Edit Medications", key="edit_medications")
            meds = case.get("medications", [])
            if editing_meds:
                meds_raw = st.text_area(
                    "Medications JSON",
                    value=json.dumps(meds, indent=2, default=str),
                    key="raw_medications",
                    height=200,
                )
                try:
                    meds_updated = json.loads(meds_raw)
                except json.JSONDecodeError:
                    st.warning("Invalid JSON")
                    meds_updated = meds
                if st.button("Save Medications", key="save_medications"):
                    try:
                        result = patch_case(case_id, {"case_data": {"medications": meds_updated}})
                        st.success("Medications saved!")
                        _load_case(result, switch_to_detail=True)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Save failed: {e}")
            else:
                for med in meds:
                    st.markdown(
                        f"- **{med.get('name', '?')}** {med.get('dose', '')} {med.get('route', '')} {med.get('frequency', '')}"
                    )

        # ── Section: Allergies (list) ─────────────────────────────────
        with st.expander("Allergies"):
            editing_allg = st.toggle("Edit Allergies", key="edit_allergies")
            allergies = case.get("allergies", [])
            if editing_allg:
                allg_raw = st.text_area(
                    "Allergies JSON",
                    value=json.dumps(allergies, indent=2, default=str),
                    key="raw_allergies",
                    height=200,
                )
                try:
                    allg_updated = json.loads(allg_raw)
                except json.JSONDecodeError:
                    st.warning("Invalid JSON")
                    allg_updated = allergies
                if st.button("Save Allergies", key="save_allergies"):
                    try:
                        result = patch_case(case_id, {"case_data": {"allergies": allg_updated}})
                        st.success("Allergies saved!")
                        _load_case(result, switch_to_detail=True)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Save failed: {e}")
            else:
                for a in allergies:
                    st.markdown(f"- **{a.get('substance', '?')}** — {a.get('reaction', '')} ({a.get('severity', '')})")

        # ── Section: Diagnostics ──────────────────────────────────────
        with st.expander("Diagnostics"):
            editing_diag = st.toggle("Edit Diagnostics", key="edit_diagnostics")
            diag = case.get("diagnostics")
            if editing_diag:
                diag_raw = st.text_area(
                    "Diagnostics JSON",
                    value=json.dumps(diag, indent=2, default=str) if diag else "{}",
                    key="raw_diagnostics",
                    height=300,
                )
                try:
                    diag_updated = json.loads(diag_raw)
                except json.JSONDecodeError:
                    st.warning("Invalid JSON")
                    diag_updated = diag
                if st.button("Save Diagnostics", key="save_diagnostics"):
                    try:
                        result = patch_case(case_id, {"case_data": {"diagnostics": diag_updated}})
                        st.success("Diagnostics saved!")
                        _load_case(result, switch_to_detail=True)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Save failed: {e}")
            else:
                st.json(diag)

        # ── Section: Plan ─────────────────────────────────────────────
        with st.expander("Plan"):
            editing_plan = st.toggle("Edit Plan", key="edit_plan")
            plan = case.get("plan")
            if editing_plan:
                plan_raw = st.text_area(
                    "Plan JSON",
                    value=json.dumps(plan, indent=2, default=str) if plan else "{}",
                    key="raw_plan",
                    height=200,
                )
                try:
                    plan_updated = json.loads(plan_raw)
                except json.JSONDecodeError:
                    st.warning("Invalid JSON")
                    plan_updated = plan
                if st.button("Save Plan", key="save_plan"):
                    try:
                        result = patch_case(case_id, {"case_data": {"plan": plan_updated}})
                        st.success("Plan saved!")
                        _load_case(result, switch_to_detail=True)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Save failed: {e}")
            else:
                st.json(plan)

        # ── Delete ────────────────────────────────────────────────────
        st.divider()
        confirm_delete = st.checkbox("I want to delete this case", key="confirm_delete")
        if confirm_delete:
            if st.button("Delete Case", type="secondary"):
                try:
                    delete_case(case_id)
                    st.success("Case deleted.")
                    del st.session_state["detail_case"]
                    st.session_state["active_tab"] = TAB_SEARCH
                    st.rerun()
                except Exception as e:
                    st.error(f"Delete failed: {e}")