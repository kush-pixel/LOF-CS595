"""Streamlit app for voice-based patient interview practice."""

from __future__ import annotations

import base64
import io
import os
import json
from datetime import datetime, timezone
from pathlib import Path
import random
import uuid
import subprocess
import shutil

from docx import Document
from dotenv import load_dotenv
import openai
import streamlit as st
import streamlit.components.v1 as components

from api_client import get_case_by_number, save_transcript

load_dotenv(Path(__file__).resolve().parent.parent / ".env")

st.set_page_config(page_title="Patient Interview", layout="wide")

OPENAI_AUDIO_MODEL = os.getenv("OPENAI_AUDIO_MODEL", "gpt-4o-audio-preview")
VOICES = ["alloy", "ash", "ballad", "coral", "echo", "fable", "onyx", "nova", "sage", "shimmer"]
PIPER_MODEL_PATH = os.getenv("PIPER_MODEL_PATH", "voices/en_US-joe-medium.onnx").strip()

EVAL_APP_URL = os.getenv("EVAL_APP_URL", "http://localhost:8503")
PREFILL_DIR = Path(__file__).resolve().parent / ".eval_prefill"
SHARED_PREFILL_FILE = PREFILL_DIR / "latest_eval_payload.json"

# ── Helpers ──────────────────────────────────────────────────────────────────


def _build_patient_system_prompt(case: dict) -> str:
    """Build a system prompt that instructs the AI to role-play as the patient."""
    demo = case.get("demographics") or {}
    hpi = case.get("chief_complaint_hpi") or {}
    social = case.get("social_history") or {}
    meds = case.get("medications") or []
    allergies = case.get("allergies") or []
    vitals = case.get("vitals") or {}
    pmh = case.get("past_medical_history") or {}

    age = demo.get("age", "unknown")
    sex = demo.get("sex", "unknown")
    language = demo.get("preferred_language", "English")

    med_list = ", ".join(
        f"{m.get('name', '?')} {m.get('dose', '')}".strip() for m in meds
    ) or "none"
    allergy_list = ", ".join(
        f"{a.get('substance', '?')} ({a.get('reaction', 'unknown reaction')})" for a in allergies
    ) or "none"
    conditions = ", ".join(pmh.get("conditions", [])) or "none"

    return f"""You are a patient in a medical interview. Stay in character at all times.

PERSONA:
- Age: {age}, Sex: {sex}, Preferred language: {language}
- You speak like a regular patient, not a clinician. Use everyday language.

YOUR SYMPTOMS (what you know and can share):
- Chief complaint: {hpi.get('chief_complaint', 'not specified')}
- Story: {hpi.get('hpi_narrative', 'not specified')}
- Onset: {hpi.get('onset', 'not sure')}
- Duration: {hpi.get('duration', 'not sure')}
- Severity: {hpi.get('severity', 'not sure')}
- What makes it worse: {', '.join(hpi.get('aggravating_factors', [])) or 'nothing specific'}
- What makes it better: {', '.join(hpi.get('alleviating_factors', [])) or 'nothing specific'}
- Other symptoms: {', '.join(hpi.get('associated_symptoms', [])) or 'none'}

YOUR BACKGROUND (reveal when asked):
- Medical conditions: {conditions}
- Current medications: {med_list}
- Allergies: {allergy_list}
- Tobacco: {social.get('tobacco', 'not specified')}
- Alcohol: {social.get('alcohol', 'not specified')}
- Drugs: {social.get('drugs', 'not specified')}
- Occupation: {social.get('occupation', 'not specified')}
- Living situation: {social.get('living_situation', 'not specified')}

YOUR VITALS (you may mention how you feel, but do NOT cite exact numbers):
- Pain level: {vitals.get('pain_scale', 'not specified')}/10

RULES:
- You do NOT know your diagnosis. You only know your symptoms and history.
- Answer questions naturally and conversationally as a real patient would.
- If the student asks something you don't know, say so naturally.
- Keep answers concise — 1-3 sentences unless the question requires more detail.
- Do NOT volunteer all information at once; let the student ask."""


def _transcribe_audio(audio_bytes: bytes) -> str:
    """Transcribe user audio via OpenAI Whisper. Returns text or fallback string."""
    print(f"Transcribing audio of length {len(audio_bytes)} bytes")
    client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY", ""))
    audio_file = io.BytesIO(audio_bytes)
    audio_file.name = "question.wav"
    result = client.audio.transcriptions.create(model="whisper-1", file=audio_file)
    return result.text.strip() or "(inaudible)"


def _piper_tts(text: str) -> tuple[bytes, str]:
    """
    Convert text -> wav bytes using Piper CLI.
    Returns (wav_bytes, wav_file_path).
    Robust for Streamlit: ensures piper is found + passes env + shows stderr.
    """
    if not text.strip():
        return b"", ""

    model_path = Path(PIPER_MODEL_PATH)
    if not model_path.exists():
        raise RuntimeError(f"Piper model not found: {model_path}")

    config_path = Path(str(model_path) + ".json")
    if not config_path.exists():
        raise RuntimeError(f"Missing Piper config: {config_path} (needs to sit next to .onnx)")

    piper_exe = shutil.which("piper")
    if not piper_exe:
        raise RuntimeError("`piper` not found in PATH. Check your venv / shell PATH.")

    out_dir = Path("generated_audio")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_wav = out_dir / f"patient_reply_{uuid.uuid4().hex}.wav"

    cmd = [piper_exe, "--model", str(model_path), "--output_file", str(out_wav)]
    env = os.environ.copy()

    proc = subprocess.run(
        cmd,
        input=text.encode("utf-8"),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.PIPE,
        env=env,
    )

    if proc.returncode != 0:
        err = proc.stderr.decode("utf-8", errors="replace")
        raise RuntimeError(f"Piper failed (exit {proc.returncode}):\n{err}")

    wav_bytes = out_wav.read_bytes()
    return wav_bytes, str(out_wav)


def _get_ai_response(messages: list, provider: str, model: str, voice: str) -> tuple[str, bytes, str]:
    """Return a text reply from the selected LLM provider."""
    if provider == "ollama":
        client = openai.OpenAI(api_key="ollama", base_url="http://localhost:11434/v1")

        safe_messages = []
        for m in messages:
            content = m.get("content")
            if isinstance(content, list):
                content = ""
            elif content is None:
                content = ""
            safe_messages.append({"role": m.get("role", "user"), "content": content})

        completion = client.chat.completions.create(model=model, messages=safe_messages)

        msg = completion.choices[0].message
        text = (getattr(msg, "content", "") or "").strip()
        wav_bytes = b""
        audio_ref = ""
        if text:
            wav_bytes, audio_ref = _piper_tts(text)
        return text, wav_bytes, audio_ref

    client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY", ""))
    completion = client.chat.completions.create(
        model=OPENAI_AUDIO_MODEL,
        modalities=["text", "audio"],
        audio={"voice": voice, "format": "wav"},
        messages=messages,
    )
    choice = completion.choices[0].message
    transcript = choice.audio.transcript
    wav_bytes = base64.b64decode(choice.audio.data)
    audio_id = choice.audio.id
    return transcript, wav_bytes, audio_id


def _build_transcript_docx(case: dict, history: list[dict]) -> bytes:
    """Build a Word document from the interview history and return raw bytes."""
    doc = Document()
    doc.add_heading(case.get("case_title", "Untitled Case"), level=1)
    doc.add_paragraph(
        f"Case #{case.get('case_number', '—')} | "
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    doc.add_paragraph("")

    for entry in history:
        label = "Student" if entry["role"] == "user" else "Patient"
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(entry["text"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _format_transcript_for_eval(history: list[dict], case_number: int, conversation_id: str | None) -> dict:
    items = []
    for e in history:
        items.append({
            "ts": e.get("ts"),
            "role": "Student" if e["role"] == "user" else "Patient",
            "text": e.get("text", ""),
        })
    return {
        "conversation_id": conversation_id,
        "case_number": case_number,
        "items": items,
    }


def _build_eval_case_payload(case: dict) -> dict:
    demo = case.get("demographics") or {}
    hpi = case.get("chief_complaint_hpi") or {}
    pmh = case.get("past_medical_history") or {}
    social = case.get("social_history") or {}
    ros = case.get("review_of_systems") or {}
    pe = case.get("physical_exam") or {}
    labs = case.get("labs") or {}
    meds = case.get("medications") or []
    allergies = case.get("allergies") or []
    imaging = case.get("imaging") or []
    ddx = case.get("differential_diagnosis") or []

    return {
        "demographics": {
            "age": demo.get("age"),
            "sex": demo.get("sex"),
        },
        "chief_complaint": hpi.get("chief_complaint", ""),
        "hpi": hpi.get("hpi_narrative", ""),
        "pmh": pmh.get("conditions", []) if isinstance(pmh.get("conditions"), list) else [],
        "medications": [
            f"{m.get('name', '')} {m.get('dose', '')}".strip()
            for m in meds
            if isinstance(m, dict)
        ],
        "allergies": [
            f"{a.get('substance', '')} ({a.get('reaction', 'unknown reaction')})".strip()
            for a in allergies
            if isinstance(a, dict)
        ],
        "social_history": {
            "narrative": "\n".join(
                [
                    f"Tobacco: {social.get('tobacco', 'not specified')}",
                    f"Alcohol: {social.get('alcohol', 'not specified')}",
                    f"Drugs: {social.get('drugs', 'not specified')}",
                    f"Occupation: {social.get('occupation', 'not specified')}",
                    f"Living situation: {social.get('living_situation', 'not specified')}",
                ]
            )
        },
        "family_history": pmh.get("family_history", []) if isinstance(pmh.get("family_history"), list) else [],
        "ros": {"narrative": ros.get("narrative", "") if isinstance(ros, dict) else ""},
        "physical_exam_findings": {"narrative": pe.get("narrative", "") if isinstance(pe, dict) else ""},
        "labs": {"narrative": labs.get("narrative", "") if isinstance(labs, dict) else ""},
        "imaging": imaging if isinstance(imaging, list) else [],
        "differential_diagnosis": ddx if isinstance(ddx, list) else [],
        "final_diagnosis": case.get("final_diagnosis", ""),
        "emotional_presentation": case.get("emotional_presentation", ""),
    }


def _build_eval_transcript_payload(history: list[dict]) -> dict:
    turns = []
    for idx, entry in enumerate(history, start=1):
        turns.append({
            "turn_number": idx,
            "speaker": "Student" if entry["role"] == "user" else "Patient",
            "content": entry.get("text", ""),
        })
    return {"turns": turns}


def _write_eval_prefill_file(case: dict, history: list[dict], conversation_id: str) -> None:
    """
    Write evaluation payload to a fixed shared file so the evaluation app
    can auto-load it without needing a query param.
    """
    PREFILL_DIR.mkdir(parents=True, exist_ok=True)

    payload = {
        "meta": {
            "conversation_id": conversation_id,
            "case_number": case.get("case_number"),
            "created_at": datetime.now(timezone.utc).isoformat(),
        },
        "case_input": _build_eval_case_payload(case),
        "transcript_input": _build_eval_transcript_payload(history),
    }

    SHARED_PREFILL_FILE.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def _redirect_to_evaluation_app() -> None:
    """
    Redirect to the evaluation app main URL only.
    No query params are used.
    """
    eval_url = EVAL_APP_URL
    print(f"Redirecting to evaluation app with URL: {eval_url}")
    components.html(
        f"""
        <script>
            window.top.location.href = {json.dumps(eval_url)};
        </script>
        """,
        height=0,
    )

# ── Session state init ───────────────────────────────────────────────────────

if "interview_case" not in st.session_state:
    st.session_state["interview_case"] = None
if "interview_messages" not in st.session_state:
    st.session_state["interview_messages"] = []
if "interview_voice" not in st.session_state:
    st.session_state["interview_voice"] = None
if "interview_history" not in st.session_state:
    st.session_state["interview_history"] = []
if "interview_conversation_id" not in st.session_state:
    st.session_state["interview_conversation_id"] = None
if "selected_llm_provider" not in st.session_state:
    st.session_state["selected_llm_provider"] = "openai"
if "selected_model" not in st.session_state:
    st.session_state["selected_model"] = "gpt-4o"
if "interview_locked_provider" not in st.session_state:
    st.session_state["interview_locked_provider"] = None
if "interview_locked_model" not in st.session_state:
    st.session_state["interview_locked_model"] = None
if "pending_audio_bytes" not in st.session_state:
    st.session_state["pending_audio_bytes"] = None
if "pending_text" not in st.session_state:
    st.session_state["pending_text"] = ""
if "pending_audio_b64" not in st.session_state:
    st.session_state["pending_audio_b64"] = ""
if "interview_ended" not in st.session_state:
    st.session_state["interview_ended"] = False
if "confirm_end" not in st.session_state:
    st.session_state["confirm_end"] = False
if "audio_turn" not in st.session_state:
    st.session_state["audio_turn"] = 0
if "final_transcript" not in st.session_state:
    st.session_state["final_transcript"] = None
if "evaluation_handoff_done" not in st.session_state:
    st.session_state["evaluation_handoff_done"] = False

# ── Sidebar ──────────────────────────────────────────────────────────────────

with st.sidebar:
    st.header("Patient Interview")

    MODEL_CHOICES = ["openai-remote", "llama3.2", "gemma3"]
    selected_model = st.selectbox(
        "Model",
        MODEL_CHOICES,
        index=MODEL_CHOICES.index(st.session_state.get("selected_model", "openai-remote"))
        if st.session_state.get("selected_model") in MODEL_CHOICES
        else 0,
        disabled=st.session_state["interview_case"] is not None,
    )
    st.session_state["selected_model"] = selected_model

    if selected_model == "openai-remote":
        st.session_state["selected_llm_provider"] = "openai"
        st.session_state["_resolved_model"] = OPENAI_AUDIO_MODEL
    else:
        st.session_state["selected_llm_provider"] = "ollama"
        st.session_state["_resolved_model"] = selected_model

    case_num = st.number_input("Case number", min_value=1, step=1, key="iv_case_num")

    if st.button("Load Case", type="primary"):
        try:
            case = get_case_by_number(case_num)
            voice = random.choice(VOICES)
            st.session_state["interview_case"] = case
            st.session_state["interview_voice"] = voice
            st.session_state["interview_conversation_id"] = str(uuid.uuid4())

            if selected_model == "openai-remote":
                locked_provider = "openai"
                locked_model = OPENAI_AUDIO_MODEL
            else:
                locked_provider = "ollama"
                locked_model = selected_model

            st.session_state["interview_locked_provider"] = locked_provider
            st.session_state["interview_locked_model"] = locked_model

            st.session_state["interview_messages"] = [
                {"role": "system", "content": _build_patient_system_prompt(case)},
            ]
            st.session_state["interview_history"] = []
            st.session_state["interview_ended"] = False
            st.session_state["confirm_end"] = False
            st.session_state["audio_turn"] = 0
            st.session_state["final_transcript"] = None
            st.session_state["evaluation_handoff_done"] = False

            st.success(f"Loaded: {case.get('case_title', 'Untitled')}")
            st.rerun()
        except Exception as e:
            st.error(f"Failed to load case #{case_num}: {e}")

    if st.session_state["interview_voice"]:
        st.caption(f"Voice: {st.session_state['interview_voice']}")

    if st.session_state["interview_history"]:
        docx_bytes = _build_transcript_docx(
            st.session_state["interview_case"],
            st.session_state["interview_history"],
        )
        st.download_button(
            label="Download Transcript (.docx)",
            data=docx_bytes,
            file_name=f"transcript_case_{st.session_state['interview_case'].get('case_number', 'unknown')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if st.session_state["interview_case"] and st.button("Reset Interview"):
        st.session_state["interview_case"] = None
        st.session_state["interview_messages"] = []
        st.session_state["interview_voice"] = None
        st.session_state["interview_history"] = []
        st.session_state["interview_conversation_id"] = None
        st.session_state["interview_locked_provider"] = None
        st.session_state["interview_locked_model"] = None
        st.session_state["interview_ended"] = False
        st.session_state["confirm_end"] = False
        st.session_state["audio_turn"] = 0
        st.session_state["final_transcript"] = None
        st.session_state["evaluation_handoff_done"] = False
        st.rerun()

    st.divider()
    st.subheader("Simulation Controls")

    end_disabled = (st.session_state["interview_case"] is None) or st.session_state["interview_ended"]

    if st.button("End Simulation", type="secondary", disabled=end_disabled):
        st.session_state["confirm_end"] = True

    if st.session_state.get("confirm_end"):
        st.warning("End the interview? This will lock the session and save the transcript.")
        col_a, col_b = st.columns(2)
        with col_a:
            if st.button("Yes, end now", type="primary"):
                st.session_state["interview_ended"] = True
                st.session_state["confirm_end"] = False
                st.rerun()
        with col_b:
            if st.button("Cancel"):
                st.session_state["confirm_end"] = False
                st.rerun()

# ── Main area ────────────────────────────────────────────────────────────────

case = st.session_state["interview_case"]

if not case:
    st.title("Patient Interview Practice")
    st.info("Enter a case number in the sidebar and click **Load Case** to begin.")
    st.stop()

if st.session_state["interview_ended"] and st.session_state["final_transcript"] is None:
    final = _format_transcript_for_eval(
        st.session_state["interview_history"],
        case.get("case_number", 0),
        st.session_state.get("interview_conversation_id"),
    )
    st.session_state["final_transcript"] = final

    try:
        save_transcript(
            conversation_id=final["conversation_id"],
            case_number=final["case_number"],
            transcript=final["items"],
        )
    except Exception as e:
        print("Auto-save failed:", e)

if st.session_state["interview_ended"]:
    st.success("Simulation ended. Transcript captured automatically.")
    st.json(st.session_state["final_transcript"])

    col1, col2 = st.columns(2)

    with col1:
        run_eval_clicked = st.button(
            "Run Evaluation",
            type="primary",
            use_container_width=True,
            disabled=st.session_state.get("evaluation_handoff_done", False),
            key="run_evaluation_btn",
        )

        if run_eval_clicked and not st.session_state.get("evaluation_handoff_done", False):
            _write_eval_prefill_file(
                case=st.session_state["interview_case"],
                history=st.session_state["interview_history"],
                conversation_id=st.session_state["interview_conversation_id"],
            )
            st.session_state["evaluation_handoff_done"] = True
            st.rerun()
        
        if st.session_state.get("evaluation_handoff_done"):
            st.success(
                "Data has been copied. Please go through further steps of how to execute evaluation frontend in the GitBook."
            )

    with col2:
        transcript_docx = _build_transcript_docx(
            st.session_state["interview_case"],
            st.session_state["interview_history"],
        )
        st.download_button(
            "Download Transcript (.docx)",
            data=transcript_docx,
            file_name=f"transcript_case_{st.session_state['interview_case'].get('case_number', 'unknown')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    st.stop()

st.title(case.get("case_title", "Untitled Case"))
demo = case.get("demographics") or {}
st.caption(
    f"Case #{case.get('case_number', '—')} | "
    f"{demo.get('age', '?')}yo {demo.get('sex', '?')} | "
    f"Specialty: {case.get('specialty', '')} | "
    f"Difficulty: {case.get('difficulty', '')}"
)

st.divider()

for entry in st.session_state["interview_history"]:
    role_label = "You" if entry["role"] == "user" else "Patient"
    with st.chat_message("user" if entry["role"] == "user" else "assistant"):
        st.markdown(f"**{role_label}:** {entry['text']}")
        if entry.get("audio_bytes"):
            st.audio(entry["audio_bytes"], format="audio/wav")

audio_input = st.audio_input("Record your question", key=f"audio_input_{st.session_state['audio_turn']}")

if audio_input is not None:
    audio_id_key = f"{audio_input.file_id}"
    if st.session_state.get("_last_audio_id") != audio_id_key:
        st.session_state["_last_audio_id"] = audio_id_key

        audio_bytes = audio_input.read()
        st.session_state["pending_audio_bytes"] = audio_bytes
        st.session_state["pending_audio_b64"] = base64.b64encode(audio_bytes).decode("utf-8")

        with st.spinner("Transcribing your question..."):
            try:
                st.session_state["pending_text"] = _transcribe_audio(audio_bytes)
            except Exception:
                st.session_state["pending_text"] = "(audio question)"

if st.session_state.get("pending_audio_bytes") is not None:
    edited_text = st.text_area(
        "Transcript",
        value=st.session_state.get("pending_text", ""),
        height=120,
        key="pending_text_area",
    )

    col_send, col_clear = st.columns([1, 1])

    with col_send:
        submit = st.button("Submit question", type="primary")

    if submit:
        provider = st.session_state["interview_locked_provider"]
        model = st.session_state["interview_locked_model"]
        user_text = edited_text.strip() or "(empty question)"

        if provider == "openai":
            st.session_state["interview_messages"].append({
                "role": "user",
                "content": [
                    {
                        "type": "input_audio",
                        "input_audio": {"data": st.session_state["pending_audio_b64"], "format": "wav"},
                    }
                ],
            })
        else:
            st.session_state["interview_messages"].append({
                "role": "user",
                "content": user_text,
            })

        st.session_state["interview_history"].append({
            "role": "user",
            "text": user_text,
            "audio_bytes": st.session_state["pending_audio_bytes"],
            "ts": datetime.now(timezone.utc).isoformat(),
        })

        with st.spinner("Patient is responding..."):
            transcript, wav_bytes, resp_audio_id = _get_ai_response(
                st.session_state["interview_messages"],
                provider,
                model,
                st.session_state["interview_voice"],
            )

        if provider == "openai":
            st.session_state["interview_messages"].append({
                "role": "assistant",
                "audio": {"id": resp_audio_id},
            })
        else:
            st.session_state["interview_messages"].append({
                "role": "assistant",
                "content": transcript,
            })

        st.session_state["interview_history"].append({
            "role": "assistant",
            "text": transcript,
            "audio_bytes": wav_bytes,
            "ts": datetime.now(timezone.utc).isoformat(),
        })

        st.session_state["pending_audio_bytes"] = None
        st.session_state["pending_audio_b64"] = ""
        st.session_state["pending_text"] = ""

        st.session_state["audio_turn"] += 1
        if st.session_state["interview_ended"]:
            st.info("Simulation has ended. Recording is disabled.")
            st.stop()
        st.rerun()